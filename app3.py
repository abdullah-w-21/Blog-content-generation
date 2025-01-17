__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import streamlit as st
import pandas as pd
from docx import Document
import json
import os
from crewai import Agent, Task, Crew, Process
from textwrap import dedent
from langchain_openai import ChatOpenAI
from datetime import datetime
import zipfile
import tempfile
import pandas as pd
from io import BytesIO, StringIO
import traceback
from datetime import datetime
import os
import os
import sys
import pandas as pd
from docx import Document
import re
from docx2python import docx2python
import markdown2
import json


LOG_FILENAME = "processing_log.txt"

def log_message(message):
    """Append a message to the log file."""
    with open(LOG_FILENAME, "a", encoding="utf-8") as f:
        f.write(message + "\n")

def log_and_terminate(message):
    """Log an error message and terminate the script."""
    log_message("ERROR: " + message)
    sys.exit(1)

def read_csv_data(csv_file, cell_mapping):
    """
    Reads arbitrary cells from an Excel sheet according to the cell_mapping dictionary.
    If the file cannot be opened or sheet not found, terminate.
    If any cell is empty, terminate.
    Note: using 1-based indexing as per previous user requirements.
    """
    try:
        df = pd.read_csv(csv_file, header=None)
    except Exception as e:
        log_and_terminate(f"Failed to open CSV file '{csv_file}': {e}")

    data = {}
    for key, (r, c) in cell_mapping.items():
        # Convert to 0-based indexing for df
        row_idx = r - 1
        col_idx = c - 1
        if row_idx >= df.shape[0] or col_idx >= df.shape[1]:
            log_and_terminate(f"Cell for '{key}' is out of range in Excel file.")

        val = df.iloc[row_idx, col_idx]
        if pd.isna(val) or str(val).strip() == "":
            log_and_terminate(f"Cell for '{key}' at row {r}, col {c} is empty. Cannot proceed.")
        data[key] = str(val).strip()

    log_message(f"Successfully read data from CSV file '{csv_file}'")
    return data

def find_nth_occurrence(content, substring, n):
    """
    Find the index of the nth occurrence of substring in content.
    Returns -1 if not found.
    """
    start = 0
    count = 0
    while True:
        index = content.find(substring, start)
        if index == -1:
            return -1
        count += 1
        if count == n:
            return index
        start = index + len(substring)

def replace_segment_in_html(html_content, search_start_marker, search_end_marker, replacement_text,
                            start_occurrence=1, end_occurrence=1):
    """
    Finds the nth occurrence of search_start_marker and nth occurrence of search_end_marker in html_content
    and replaces everything from search_start_marker to search_end_marker (inclusive) with replacement_text.
    If occurrences are not found, terminate.
    """
    start_index = find_nth_occurrence(html_content, search_start_marker, start_occurrence)
    if start_index == -1:
        log_and_terminate(f"Search start marker '{search_start_marker}' (occurrence {start_occurrence}) not found in HTML.")

    end_index = find_nth_occurrence(html_content[start_index+len(search_start_marker):], search_end_marker, end_occurrence)
    if end_index == -1:
        log_and_terminate(f"Search end marker '{search_end_marker}' (occurrence {end_occurrence}) not found in HTML after start marker '{search_start_marker}'.")

    # Adjust end_index relative to the whole string
    end_index = start_index + len(search_start_marker) + end_index

    before = html_content[:start_index]
    after = html_content[end_index + len(search_end_marker):]
    return before + replacement_text + after

def load_html_template(file_path):
    """Load the HTML template file. If fails, terminate."""
    if not os.path.isfile(file_path):
        log_and_terminate(f"HTML template file '{file_path}' does not exist.")

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        log_message(f"Successfully loaded HTML template '{file_path}'.")
        return content
    except Exception as e:
        log_and_terminate(f"Failed to read HTML template '{file_path}': {e}")

def save_html_content(file_path, content):
    """Save the updated HTML content to a file."""
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        log_message(f"Successfully saved updated HTML to '{file_path}'.")
    except Exception as e:
        log_and_terminate(f"Failed to write updated HTML to '{file_path}': {e}")

def update_html_content(html_content, data, replacements_config):
    """
    Given html_content, data dictionary, and replacements_config, perform all required replacements.

    Each entry in replacements_config should provide:
    - search_start_marker (string): marker that defines start of the section to remove
    - search_end_marker (string): marker that defines end of the section to remove
    - search_start_occurrence (int, optional): which occurrence of start marker
    - search_end_occurrence (int, optional): which occurrence of end marker
    - replacement_start (string, optional): text to put before cell content in the replacement
    - replacement_end (string, optional): text to put after cell content in the replacement

    If replacement_start/end are not provided, defaults to empty strings.
    """
    updated_content = html_content

    for key, config in replacements_config.items():
        cell_value = data.get(key, "")

        # Required keys for searching in HTML
        search_start_marker = config["search_start_marker"]
        search_end_marker = config["search_end_marker"]

        # Optional occurrence parameters
        start_occ = config.get("search_start_occurrence", 1)
        end_occ = config.get("search_end_occurrence", 1)

        # Optional replacement wrappers
        replacement_start = config.get("replacement_start", "")
        replacement_end = config.get("replacement_end", "")

        # Build the replacement text
        replacement_text = replacement_start + cell_value + replacement_end

        updated_content = replace_segment_in_html(
            updated_content,
            search_start_marker,
            search_end_marker,
            replacement_text,
            start_occurrence=start_occ,
            end_occurrence=end_occ
        )

        log_message(f"Successfully replaced segment for '{key}' with '{replacement_text}'.")
    return updated_content

def process_html_template(html_file, excel_file, sheet_name, cell_mapping, replacements_config, in_place=True):
    """
    Process a single HTML template and update based on Excel data.
    The script will terminate on any error.
    """
    data = read_excel_data(excel_file, sheet_name, cell_mapping)
    html_content = load_html_template(html_file)
    updated_content = update_html_content(html_content, data, replacements_config)

    if in_place:
        output_file = html_file
    else:
        base, ext = os.path.splitext(html_file)
        output_file = base + "-new" + ext

    save_html_content(output_file, updated_content)
    log_message("All operations completed successfully.")

def docx_to_html_with_docx2python(
    docx_path,
    entire_start="",
    entire_end="",
    p_start="<p>",
    p_end="</p>"
):
    """
    Converts a DOCX file into HTML with:
    - Detection and handling of nested bullet and numbered lists.
    - Optionally wraps entire content and paragraphs with specified tags.
    - No inline formatting (e.g., bold) is processed here.

    Args:
        docx_path (str): Path to the DOCX file.
        entire_start (str): HTML tag or string to wrap the entire content at the start.
        entire_end (str): HTML tag or string to wrap the entire content at the end.
        p_start (str): HTML tag to wrap each paragraph at the start.
        p_end (str): HTML tag to wrap each paragraph at the end.

    Returns:
        str: Generated HTML content.
    """
    doc_content = docx2python(docx_path)
    paragraphs = []

    # Extract non-empty paragraphs
    for section in doc_content.body:
        for page in section:
            for column in page:
                for paragraph in column:
                    if paragraph.strip():
                        paragraphs.append(paragraph)

    html_fragments = []
    if entire_start:
        html_fragments.append(entire_start)

    # This stack will track open lists. Each element is ("ul" or "ol").
    list_stack = []

    def close_all_lists():
        """Close all currently open lists."""
        while list_stack:
            tag = list_stack.pop()
            html_fragments.append(f"</{tag}>")

    def close_to_level(level):
        """
        Close lists until the nesting depth (len(list_stack)) equals 'level'.
        'level' here is the number of open lists desired.
        """
        while len(list_stack) > level:
            tag = list_stack.pop()
            html_fragments.append(f"</{tag}>")

    def open_list(list_type):
        """Open a new list of the given type (ul or ol)."""
        list_stack.append(list_type)
        html_fragments.append(f"<{list_type}>")

    # Function to detect and handle list paragraphs
    # We consider indent_level = number of leading tabs at line start.
    # For bullets: '--\t' at start (after indentation) indicates a bullet item.
    # For nested bullets: one extra leading tab per nesting level.
    # Numbered lists: a regex for leading digits, e.g., "1. " or "1) "
    for paragraph in paragraphs:
        # Count leading tabs
        indent_match = re.match(r"^(\t+)", paragraph)
        indent_level = len(indent_match.group(1)) if indent_match else 0

        # Line after removing leading tabs
        line_stripped = paragraph[indent_level:]

        # Detect bullet or number
        bullet_match = re.match(r"^--\t", line_stripped)
        number_match = re.match(r"^[0-9]+[\.\)]\s?", line_stripped)

        if bullet_match:
            current_type = "ul"
        elif number_match:
            current_type = "ol"
        else:
            current_type = None

        if current_type:
            # This is a list item
            # Desired depth: indent_level + 1 means if indent_level=0, we want 1 open list, if=1, we want 2, etc.
            desired_depth = indent_level + 1

            # Adjust the current list stack depth
            if len(list_stack) > desired_depth:
                # Close lists until we are at the correct depth
                close_to_level(desired_depth)
            elif len(list_stack) < desired_depth:
                # Need to open more lists
                while len(list_stack) < desired_depth:
                    open_list(current_type)
            else:
                # Same depth, check if the current top matches the current_type
                if list_stack[-1] != current_type:
                    # Close the mismatched list and open the correct one
                    close_to_level(desired_depth - 1)
                    open_list(current_type)

            # Remove the bullet or number marker from the content
            if bullet_match:
                # Remove the '--\t' marker
                item_content = re.sub(r"^--\t", "", line_stripped)
            else:
                # Remove the numbering marker (e.g. '1. ', '1) ')
                item_content = re.sub(r"^[0-9]+[\.\)]\s?", "", line_stripped)

            # Escape HTML special chars in content
            item_content = (item_content
                            .replace("&", "&amp;")
                            .replace("<", "&lt;")
                            .replace(">", "&gt;"))

            html_fragments.append(f"<li>{item_content}</li>")

        else:
            # Not a list item, close all lists and add a paragraph
            close_all_lists()
            paragraph_content = (paragraph
                                 .replace("&", "&amp;")
                                 .replace("<", "&lt;")
                                 .replace(">", "&gt;"))
            html_fragments.append(f"{p_start}{paragraph_content}{p_end}")

    # Close any remaining lists
    close_all_lists()

    if entire_end:
        html_fragments.append(entire_end)

    return "\n".join(html_fragments)



def replace_all_markers_in_html(html_content, markers):
    """
    Replace all occurrences of marker pairs in the HTML content with corresponding tags.

    Args:
        html_content (str): The HTML content where markers need to be replaced.
        markers (dict): Dictionary of markers in the format:
            {
                (start_marker, end_marker): (start_tag, end_tag),
                ...
            }

    Returns:
        str: Updated HTML content with all marker pairs replaced.
    """
    updated_html = html_content

    for (start_marker, end_marker), (start_tag, end_tag) in markers.items():
        while True:
            start_idx = updated_html.find(start_marker)
            if start_idx == -1:
                # No more start markers; move to the next marker pair
                break
            
            # Look for the next end marker after the start marker
            end_idx = updated_html.find(end_marker, start_idx + len(start_marker))
            if end_idx == -1:
                # No matching end marker; stop processing this pair
                log_message(f"Unmatched start marker '{start_marker}' found without corresponding end marker '{end_marker}'.")
                break

            # Extract the content between start and end markers
            inner_content = updated_html[start_idx + len(start_marker):end_idx]

            # Build the replacement with the tags
            replacement = f"{start_tag}{inner_content}{end_tag}"

            # Replace the segment (including start and end markers) with the replacement
            updated_html = (
                updated_html[:start_idx] +
                replacement +
                updated_html[end_idx + len(end_marker):]
            )

    return updated_html


def process_preliminary_html(preliminary_html, markers):
    """
    Convert preliminary HTML into processed HTML by replacing markers with corresponding tags.

    Args:
        preliminary_html (str): The raw HTML generated from the DOCX file.
        markers (dict): Dictionary of markers in the format:
            {
                (start_marker, end_marker): (start_tag, end_tag),
                ...
            }

    Returns:
        str: Processed HTML with all markers replaced.
    """
    log_message("Starting to process preliminary HTML with markers...")
    processed_html = replace_all_markers_in_html(preliminary_html, markers)
    log_message("Successfully processed preliminary HTML.")
    return processed_html

def preprocess_bold_text(docx_path, output_path):
    """
    Preprocess a DOCX file to encapsulate bold text with <bold> and </bold> tags.

    Args:
        docx_path (str): Path to the input DOCX file.
        output_path (str): Path to save the modified DOCX file.

    Returns:
        None
    """
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        new_runs = []
        for run in paragraph.runs:
            if run.bold:  # If the text is bold
                run.text = f"<bold>{run.text}</bold>"
                run.bold = False  # Remove the bold formatting
            new_runs.append(run.text)

        # Update paragraph text with modified runs
        paragraph.text = "".join(new_runs)

    # Save the modified document
    doc.save(output_path)
    print(f"Processed document saved to: {output_path}")

def markdown_to_html(markdown_text):
    """
    Converts Markdown text to HTML, handling headings, bullet points, 
    numbered lists, bold text, and paragraphs (including multiple paragraphs 
    and paragraphs ending at the end of lines or within a line).

    Args:
        markdown_text: The Markdown text to convert.

    Returns:
        The HTML representation of the Markdown text.
    """
    html_lines = []
    in_list = False
    in_ordered_list = False
    in_paragraph = False

    def close_paragraph():
        nonlocal in_paragraph
        if in_paragraph:
            html_lines.append('</p>')
            in_paragraph = False

    def close_list():
        nonlocal in_list, in_ordered_list
        if in_list:
            html_lines.append('</ul>')
            in_list = False
        if in_ordered_list:
            html_lines.append('</ol>')
            in_ordered_list = False

    for line in markdown_text.splitlines():
        line = line.strip()

        # Handle headings
        if line.startswith('#'):
            close_paragraph()
            close_list()
            match = re.match(r'(#+)\s*(H\d)?(:\s*)?(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(4).strip()
                if 1 <= level <= 6:
                    html_lines.append(f'<h{level}>{text}</h{level}>')
            continue

        # Handle bullet points (unordered lists)
        if line.startswith('- '):
            close_paragraph()
            if not in_list:
                html_lines.append('<ul class="list" role="list">')
                in_list = True
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
            html_lines.append(f'  <li>{text}</li>')
            continue
        else:
            if in_list:
                html_lines.append('</ul>')
                in_list = False

        # Handle numbered lists (ordered lists)
        if re.match(r'\d+\.\s', line):
            close_paragraph()
            if not in_ordered_list:
                html_lines.append('<ol class="list" role="list">')
                in_ordered_list = True
            text = line.split('. ', 1)[1].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
            html_lines.append(f'<li>{text}</li>')
            continue
        else:
            if in_ordered_list:
                html_lines.append('</ol>')
                in_ordered_list = False

        # Handle paragraphs (with line break detection within paragraphs)
        if not in_paragraph:
            html_lines.append('<p>')
            in_paragraph = True

        # Split the line by newline characters to handle paragraph breaks within lines
        sublines = line.split('\n')  

        for subline in sublines:
            subline = subline.strip()
            if subline:
                # Apply bold formatting using regular expressions
                subline = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', subline)  

                html_lines.append(subline)
                close_paragraph()  # Close the paragraph after each subline
                if subline != sublines[-1]:  # Don't open a new paragraph if it's the last subline
                    html_lines.append('<p>')
                    in_paragraph = True

    # Close any open elements at the end
    close_paragraph()
    close_list()

    return '\n'.join(html_lines)

def format_meta_tags(meta_tags_str):
  """
  Converts a JSON string of meta tags into HTML meta tag strings.

  Args:
    meta_tags_str: A JSON string containing meta tag information.
                   The JSON object should have a "raw_schema" key containing 
                   the HTML code snippet with the meta tags.

  Returns:
    A string containing the formatted HTML meta tags with "```html" 
    and "```" removed.
  """

  try:
    meta_tags_dict = json.loads(meta_tags_str)
    html_code = meta_tags_dict["raw_schema"] 
    # Remove "```html" from the beginning and "```" from the end
    html_code = html_code.replace("```html\n", "").replace("```", "")  
    return html_code
  except (json.JSONDecodeError, KeyError) as e:
    return f"Error: {e}"
  
def generate_filled_html(content, location):
    """
    Generates an HTML file filled with content from a CSV file for a specific location.

    Args:
        content (str): The name of the CSV file containing the content.
        location (str): The location (e.g., 'astoria', 'williamsburg') for which to generate the HTML.

    Returns:
        None
    """
    
    csv_file = content
    
    html_names = ['astoria', 
                  'williamsburg', 
                  'hicksville', 
                  '174th-street', 
                  'jackson-heights', 
                  'bartow-mall', 
                  'jamaica', 
                  'stuytown', 
                  'crown-heights', 
                  'mineola', 
                  'long-island-city']
    
    if location not in html_names:
        log_and_terminate(f"Invalid location: {location}")

    html_template_file = location + '-template.html'  # Construct the template file name
    saving_name = location + '.html'
    
    
    cell_mapping = {
        "TITLE": (2,1),
        "META_DESC": (2,2),
        "FAQ_SCHEMA": (2,3),
        "CONTENT": (2,4),
        "H2": (2,6)
    }

    # 1. Read data from the CSV file using a modified read_excel_data (now read_csv_data)
    data = read_csv_data(csv_file, cell_mapping)

    # 2. Load the HTML template
    template = load_html_template(html_template_file)

    # 3. Replace content in the template
    final_html = replace_segment_in_html(template, '<title>', '</title>','<title>'+data['TITLE']+'</title>')
    final_html = replace_segment_in_html(final_html, '<meta', '"viewport"/>','<meta name="viewport" content="width=device-width, initial-scale=1"/>', start_occurrence=2)
    final_html = replace_segment_in_html(final_html, '<meta', '"description"/>','<meta name=description content="'+data['META_DESC']+'"/>', start_occurrence=3)
    data['FAQ_SCHEMA'] = format_meta_tags(data['FAQ_SCHEMA'])
    final_html = replace_segment_in_html(final_html, "<script type=", "</script>",data['FAQ_SCHEMA'])
    data['CONTENT'] = markdown_to_html(data['CONTENT'])
    final_html = replace_segment_in_html(final_html, '<div class="frame-1000004889">', '</div>', '<section class="main-content">\n'+data['CONTENT']+'\n</section>', end_occurrence=3)
    final_html = replace_segment_in_html(final_html, '<h1>', '</h1>','<h2>'+data['H2']+'</h2>', start_occurrence=1)
    final_html = replace_segment_in_html(final_html, '<h1>', '</h1>','<h2>'+data['H2']+'</h2>', start_occurrence=1)


    # 4. Save the updated HTML
    save_html_content(saving_name, final_html)
    log_message(f"Successfully generated HTML for {location} from {csv_file}")
  


#Example Usage
#content_file_name = 'content.csv'
#location = 'astoria'
#generate_filled_html(content_file_name, location)


# Session state initialization
if 'page_type' not in st.session_state:
    st.session_state.page_type = None
if 'site_name' not in st.session_state:
    st.session_state.site_name = None
if 'template_text' not in st.session_state:
    st.session_state.template_text = None
if 'template_structure' not in st.session_state:
    st.session_state.template_structure = None
if 'primary_keyword' not in st.session_state:
    st.session_state.primary_keyword = None
if 'schema_template' not in st.session_state:
    st.session_state.schema_template = None
if 'service_name' not in st.session_state:
    st.session_state.service_name = None
if 'additional_keywords_list' not in st.session_state:
    st.session_state.additional_keywords_list = None
if 'generation_content' not in st.session_state:
    st.session_state.generation_content = None




def debug_print(message):
    """Print debug messages in Streamlit"""
    st.write(f"Debug: {message}")

def check_templates():
    """Check for template files and print directory information"""
    current_dir = os.getcwd()
    debug_print(f"Current working directory: {current_dir}")
    
    all_files = os.listdir(current_dir)
    debug_print(f"All files in directory: {all_files}")
    
    template_files = [f for f in all_files if f.endswith('-template.html')]
    debug_print(f"Found template files: {template_files}")
    
    return template_files

def configure_openai():
    api_key = st.sidebar.text_input("Enter OpenAI API Key", type="password")
    if api_key:
        os.environ["OPENAI_API_KEY"] = api_key
        return True
    return False


def read_word_document(file):
    doc = Document(file)
    template_text = []
    template_structure = {'sections': []}

    current_section = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name

        if style.startswith('Heading'):
            if style.startswith('Heading 2'):
                current_section = {'heading': text, 'level': 2, 'subsections': [], 'content': []}
                template_structure['sections'].append(current_section)
            elif style.startswith('Heading 3') and current_section:
                current_section['subsections'].append({'heading': text, 'level': 3, 'content': []})
        elif current_section:
            if current_section['subsections']:
                current_section['subsections'][-1]['content'].append(text)
            else:
                current_section['content'].append(text)

        template_text.append(text)

    return '\n'.join(template_text), template_structure


class TemplateAnalyzer(Agent):
    def __init__(self):
        super().__init__(
            role='Template Analyzer',
            goal='Analyze content template structure and create detailed outline',
            backstory=dedent("""Expert in content analysis and structural pattern recognition."""),
            verbose=True,
            allow_delegation=False,
            llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.2)
        )


class ContentWriter(Agent):
    def __init__(self):
        super().__init__(
            role='Content Writer',
            goal='Write high-quality, SEO-optimized content',
            backstory=dedent("""Expert content writer with deep knowledge of SEO and EEAT framework."""),
            verbose=True,
            allow_delegation=False,
            llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.7)
        )


class SEOSpecialist(Agent):
    def __init__(self):
        super().__init__(
            role='SEO Specialist',
            goal='Optimize content for search engines and create schema markup',
            backstory=dedent("""Expert in technical SEO, content optimization, and schema markup creation."""),
            verbose=True,
            allow_delegation=False,
            llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
        )


def format_content(content, format_type='markdown'):
    lines = content.split('\n')
    formatted = []

    for line in lines:
        line = line.strip()
        if format_type == 'markdown':
            if line.startswith('H2:'):
                formatted.append(f"## {line[3:].strip()}")
            elif line.startswith('H3:'):
                formatted.append(f"### {line[3:].strip()}")
            elif line:
                formatted.append(line)
        else:
            formatted.append(line)

    return '\n'.join(formatted)


def create_word_doc(content):
    doc = Document()
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('H2:'):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('H3:'):
            doc.add_heading(line[3:].strip(), level=3)
        elif line:
            doc.add_paragraph(line)
    return doc


def generate_content(template_text, template_structure, primary_keyword, schema_template,
                     service_name=None, additional_keywords=None, progress_callbacks=None):
    try:
        template_analyzer = TemplateAnalyzer()
        writer = ContentWriter()
        seo_specialist = SEOSpecialist()

        # Template analysis task
        analysis_task = Task(
            description=f"""
            Analyze this content template and create a detailed outline:
            Template Structure: {json.dumps(template_structure)}
            Original Text: {template_text}

            Create a detailed analysis of:
            1. Content flow and exact structure
            2. Key sections and their purposes
            3. Content patterns and relationships
            4. Recommended approach for content creation
            5. The address details mentioned in the content remember it.
            6. Pay close attention to the use of headings (H1, H2, H3) and how they organize the information.
            7. The number of FAQ's
            8. FAQ's questions starts with Q: and Answer starts with A:
            """,
            expected_output="Detailed template analysis with structural insights and content recommendations",
            agent=template_analyzer
        )

        # Get template analysis
        template_analysis = template_analyzer.execute_task(analysis_task)

        # Writing task with template insights
        writing_task = Task(
            description=f"""
            Generate content using this template analysis:
            {template_analysis}

            Primary Keyword: {primary_keyword}
            Service Name: {service_name if service_name else 'Not specified'}
            Additional Keywords: {', '.join(additional_keywords) if additional_keywords else 'None'}

            Generate SEO Optimized high-quality, informative, and trustworthy content for Nao Medical, a leading healthcare provider in NYC with over 11 facilities. 
            Requirements:
            1. Follow provided template structure exactly and accurately (include the address if exists)
            2. Use H2: and H3: prefix for headings
            3. Word count: 1700-1800 so it ranks on google
            4. Follow EEAT framework
            5. Natural LSI keyword integration
            6. Prioritize user value and clarity
            7. FAQ's questions starts with (H3: Q:) and Answer starts with (A:)

            Output content using H2: and H3: prefixes for headings.
            Ensure that the output is perfect and the headings prefixes are mention (H1,H2,H3).
            """,
            expected_output="SEO-optimized content following template structure with proper heading hierarchy with (H1:,H2:,H3:) prefixes for headings. Output should only consist of the content along with their markdown.",
            agent=writer
        )

        content_result = writer.execute_task(writing_task)

        # SEO task with generated content
        seo_task = Task(
            description=f"""
            Using this content:
            {content_result}

            Create:
            1. SEO title (max 60 chars)
            2. Meta description (max 160 chars)
            3. FAQ Schema from content
            4. Adapt schema template: {schema_template}

            Primary keyword: {primary_keyword}

            Format:
            TITLE: [title]
            META: [description]
            SCHEMA:
            [complete meta tags and schema markup with FAQs schema]
            """,
            expected_output="SEO metadata including title, meta description, and complete meta tags and schema markup with FAQs schema",
            agent=seo_specialist
        )

        crew = Crew(
            agents=[template_analyzer, writer, seo_specialist],
            tasks=[analysis_task, writing_task, seo_task],
            verbose=True,
            process=Process.sequential
        )

        result = crew.kickoff()

        # Format outputs
        markdown_content = format_content(str(result.tasks_output[1]), 'markdown')
        word_content = format_content(str(result.tasks_output[1]), 'word')

        # Parse SEO output
        seo_output = str(result.tasks_output[2])
        title = meta_description = ""
        schema = {"meta_tags": [], "json_ld": []}

        current_section = None
        current_content = []

        for line in seo_output.split('\n'):
            line = line.strip()
            if line.startswith('TITLE:'):
                title = line[6:].strip()
            elif line.startswith('META:'):
                meta_description = line[5:].strip()
            elif line.startswith('SCHEMA:'):
                current_section = 'schema'
            elif current_section == 'schema':
                current_content.append(line)

        if current_content:
            try:
                schema = {"raw_schema": '\n'.join(current_content)}
            except:
                schema = {"error": "Schema processing failed"}

        return {
            'content': markdown_content,
            'word_content': word_content,
            'title': title,
            'meta_description': meta_description,
            'schema': schema,
            'template_analysis': str(result.tasks_output[0])
        }

    except Exception as e:
        if progress_callbacks and 'error' in progress_callbacks:
            progress_callbacks['error'](f"Error: {str(e)}")
        raise e


def parse_seo_output(seo_output):
    title = meta_description = ""
    schema = {"meta_tags": [], "json_ld": []}
    current_section = None
    current_content = []

    for line in seo_output.split('\n'):
        line = line.strip()
        if line.startswith('TITLE:'):
            title = line[6:].strip()
        elif line.startswith('META:'):
            meta_description = line[5:].strip()
        elif line.startswith('SCHEMA:'):
            current_section = 'schema'
        elif current_section == 'schema':
            current_content.append(line)

    if current_content:
        schema = {"raw_schema": '\n'.join(current_content)}

    return title, meta_description, schema


def regenerate_component(component_type, template_text, template_structure, primary_keyword,
                         schema_template, service_name=None, additional_keywords=None,
                         existing_content=None):
    if component_type == 'content':
        template_analyzer = TemplateAnalyzer()
        writer = ContentWriter()

        analysis_task = Task(
            description=f"""
            Analyze this content template and create a detailed outline:
            Template Structure: {json.dumps(template_structure)}
            Original Text: {template_text}

            Create a detailed analysis of:
            1. Content flow and exact structure
            2. Key sections and their purposes
            3. Content patterns and relationships
            4. Recommended approach for content creation
            5. The address details mentioned in the content remember it.
            6. Pay close attention to the use of headings (H1, H2, H3) and how they organize the information.
            7. The number of FAQ's
            8. FAQ's questions starts with Q: and Answer starts with A:
            """,
            expected_output="Detailed template analysis with structural insights and content recommendations",
            agent=template_analyzer
        )

        # Get template analysis
        template_analysis = template_analyzer.execute_task(analysis_task)

        # Writing task with template insights
        writing_task = Task(
           description=f"""
            Generate content using this template analysis:
            {template_analysis}

            Primary Keyword: {primary_keyword}
            Service Name: {service_name if service_name else 'Not specified'}
            Additional Keywords: {', '.join(additional_keywords) if additional_keywords else 'None'}

            Generate SEO Optimized high-quality, informative, and trustworthy content for Nao Medical, a leading healthcare provider in NYC with over 11 facilities. 
            Requirements:
            1. Follow provided template structure exactly and accurately (include the address if exists)
            2. Use H2: and H3: prefix for headings
            3. Word count: 1700-1800 so it ranks on google
            4. Follow EEAT framework
            5. Natural LSI keyword integration
            6. Prioritize user value and clarity
            7. FAQ's questions starts with (H3: Q:) and Answer starts with (A:)

            Output content using H2: and H3: prefixes for headings.
            Ensure that the output is perfect and the headings prefixes are mention (H1,H2,H3).
            """,
            expected_output="SEO-optimized content following template structure with proper heading hierarchy with (H1:,H2:,H3:) prefixes for headings. Output should only consist of the content along with their markdown.",
            agent=writer
        )

        content_result = writer.execute_task(writing_task)

        return {
            'content': format_content(str(content_result), 'markdown'),
            'word_content': format_content(str(content_result), 'word')
        }

    elif component_type == 'seo':
        seo_specialist = SEOSpecialist()
        seo_task = Task(
            description=f"""
                    Using this content:
                    {existing_content}

                    Create:
                    1. SEO title (max 60 chars)
                    2. Meta description (max 160 chars)
                    3. FAQ Schema from content
                    4. Adapt schema template: {schema_template}

                    Primary keyword: {primary_keyword}

                    Format:
                    TITLE: [title]
                    META: [description]
                    SCHEMA:
                    [complete meta tags and schema markup with FAQs schema]
                    """,
            expected_output="SEO metadata including title, meta description, and complete meta tags and schema markup with FAQs schema.",
            agent=seo_specialist
        )
        seo_result = seo_specialist.execute_task(seo_task)

        title, meta_description, schema = parse_seo_output(str(seo_result))
        return {
            'title': title,
            'meta_description': meta_description,
            'schema': schema
        }


def regenerate_callback(component_type):
    with st.spinner(f"Regenerating {component_type}..."):
        additional_keywords_list = ([k.strip() for k in st.session_state.additional_keywords.split('\n')
                                     if k.strip()] if st.session_state.additional_keywords else None)

        if component_type == 'content':
            result = regenerate_component('content',
                                          st.session_state.template_text,
                                          st.session_state.template_structure,
                                          st.session_state.primary_keyword,
                                          st.session_state.schema_template,
                                          st.session_state.service_name,
                                          additional_keywords_list)
            st.session_state.generation_content.update(result)
            # Update edited content as well
            st.session_state.edited_content = result.get('content', st.session_state.edited_content)
            st.rerun()  # Force rerun after update

        elif component_type == 'seo':
            result = regenerate_component('seo',
                                          st.session_state.template_text,
                                          st.session_state.template_structure,
                                          st.session_state.primary_keyword,
                                          st.session_state.schema_template,
                                          st.session_state.service_name,
                                          additional_keywords_list,
                                          st.session_state.generation_content['content'])
            st.session_state.generation_content.update(result)
            # Update edited SEO as well
            st.session_state.edited_seo.update({
                'title': result.get('title', st.session_state.edited_seo.get('title', '')),
                'meta_description': result.get('meta_description',
                                               st.session_state.edited_seo.get('meta_description', '')),
                'schema': result.get('schema', st.session_state.edited_seo.get('schema', {}))
            })
            st.rerun()  # Force rerun after update

        else:  # regenerate all
            result = generate_content(
                st.session_state.template_text,
                st.session_state.template_structure,
                st.session_state.primary_keyword,
                st.session_state.schema_template,
                st.session_state.service_name,
                additional_keywords_list)

            # Update both generation content and edited states
            st.session_state.generation_content = result
            st.session_state.edited_content = result['content']
            st.session_state.edited_seo = {
                'title': result['title'],
                'meta_description': result['meta_description'],
                'schema': result['schema']
            }
            st.rerun()  # Force rerun after update


def save_to_csv(data):
    df = pd.DataFrame([data])
    return df.to_csv(index=False).encode('utf-8')


def expand_content(selected_text, primary_keyword):
    """Function to expand selected content using GPT-4 directly"""
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7)

    prompt = f"""
    Expand and enhance this content while maintaining the same style and tone:
    {selected_text}

    Primary Keyword: {primary_keyword}

    Requirements:
    1. Maintain the same style and format
    2. Add more detailed information and examples
    3. Keep SEO optimization in mind
    4. Ensure natural flow with surrounding content
    5. Keep the same heading format if present (H1:, H2:, H3:)
    """

    response = llm.invoke(prompt)
    return response.content


def count_words(text):
    """Count words in text, excluding heading markers"""
    # Remove heading markers (H1:, H2:, H3:)
    cleaned_text = text.replace('H1:', '').replace('H2:', '').replace('H3:', '')
    return len(cleaned_text.split())


def main():
    # Initialize session states
    if 'form_submitted' not in st.session_state:
        st.session_state.form_submitted = False
    if 'template_doc_content' not in st.session_state:
        st.session_state.template_doc_content = None
    if 'edited_content' not in st.session_state:
        st.session_state.edited_content = None
    if 'edited_seo' not in st.session_state:
        st.session_state.edited_seo = {}
    if 'generation_content' not in st.session_state:
        st.session_state.generation_content = None

    st.title("Content Generation App")

    with st.expander("ℹ️ About v2.0 Updates"):
        st.markdown("""
        ### 🆕 New Features in v2.0
        
        #### Content Management
        - **Original Content View**: Always visible generated content
        - **Content Editor**: Expandable editor for making changes
        - **Word Count**: Real-time word count tracking
        - **Content Expansion**: Ability to expand specific sections using AI
        
        #### Editor Features
        - 📝 Edit content while preserving original
        - 💾 Save changes independently
        - 🔄 Auto-sync with regenerated content
        - 📊 Real-time word count updates
        
        #### SEO Tools
        - 🎯 Enhanced SEO metadata editor
        - ✍️ Title character counter (60 char limit)
        - 📑 Meta description counter (160 char limit)
        - 🔧 JSON schema validation
        
        #### Regeneration Options
        - 🔄 Regenerate content only
        - 🎯 Regenerate SEO only
        - ⚡ Regenerate everything
        
        #### File Management
        - 📦 Download all files as ZIP
        - 📄 Individual file downloads
        - 💾 Auto-save of edited versions
        
        #### How to Use
        1. Generate initial content
        2. View generated content at the top
        3. Use the Content Editor expander to make changes
        4. Expand sections using AI assistance
        5. Edit SEO metadata in the SEO expander
        6. Download your preferred file format
        
        #### Tips
        - Use the content expander for major edits
        - Save changes before regenerating
        - Monitor word count for optimal SEO
        - Validate schema before saving
        """)

    if not configure_openai():
        st.warning("Please enter your OpenAI API key in the sidebar.")
        return

    # Initial form for content generation
    if not st.session_state.form_submitted:
        with st.form("content_form"):
            template_doc = st.file_uploader("Upload template Word document", type=['docx'])
            schema_template = st.text_area("Enter Schema Template", height=200)
            col1, col2 = st.columns(2)
            with col1:
                primary_keyword = st.text_input("Primary Keyword (required)")
                service_name = st.text_input("Service Name (optional)")
            with col2:
                additional_keywords = st.text_area("Additional Keywords (optional, one per line)")
            submitted = st.form_submit_button("Generate Content ✨")

            if submitted and template_doc and primary_keyword and schema_template:
                template_text, template_structure = read_word_document(template_doc)
                st.session_state.template_doc_content = template_doc.getvalue()
                st.session_state.template_text = template_text
                st.session_state.template_structure = template_structure
                st.session_state.primary_keyword = primary_keyword
                st.session_state.schema_template = schema_template
                st.session_state.service_name = service_name
                st.session_state.additional_keywords = additional_keywords
                st.session_state.form_submitted = True
                st.rerun()

    # Content generation and editing interface
    if st.session_state.form_submitted:
        if st.button("← Back to Form"):
            st.session_state.form_submitted = False
            st.session_state.generation_content = None
            st.session_state.edited_content = None
            st.session_state.edited_seo = {}
            st.rerun()

        progress_tab, output_tab = st.tabs(["Generation Progress", "Final Output"])

        with progress_tab:
            st.subheader("🔄 Generation Progress")
            status = st.empty()
            progress_bar = st.progress(0)

            if not st.session_state.generation_content:
                try:
                    status.text("Starting generation...")
                    progress_bar.progress(10)

                    additional_keywords_list = ([k.strip() for k in st.session_state.additional_keywords.split('\n')
                                                 if k.strip()] if st.session_state.additional_keywords else None)

                    result = generate_content(
                        st.session_state.template_text,
                        st.session_state.template_structure,
                        st.session_state.primary_keyword,
                        st.session_state.schema_template,
                        st.session_state.service_name,
                        additional_keywords_list
                    )

                    # Update both generation and edited content
                    st.session_state.generation_content = result
                    st.session_state.edited_content = result['content']
                    st.session_state.edited_seo = {
                        'title': result['title'],
                        'meta_description': result['meta_description'],
                        'schema': result['schema']
                    }
                    progress_bar.progress(100)
                    status.text("Generation complete!")

                except Exception as e:
                    st.error(f"Error: {str(e)}")
                    return

        with output_tab:
            if st.session_state.generation_content:
                # Regeneration buttons
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.button("🔄 Regenerate Content", on_click=regenerate_callback, args=('content',),
                              key='btn_content')
                with col2:
                    st.button("🔄 Regenerate SEO", on_click=regenerate_callback, args=('seo',), key='btn_seo')
                with col3:
                    st.button("🔄 Regenerate All", on_click=regenerate_callback, args=('all',), key='btn_all')

                # Original Content View
                st.subheader("📄 Generated Content")
                st.markdown(st.session_state.generation_content['content'])

                # Content Editor in Expander
                with st.expander("📝 Content Editor"):
                    edited_content = st.text_area(
                        "Edit Generated Content",
                        value=st.session_state.edited_content,
                        height=500,
                        key="content_editor"
                    )

                    # Word Count Display
                    word_count = count_words(edited_content)
                    st.metric("Word Count", word_count)

                    # Save content changes
                    if st.button("Save Content Changes", key="save_content_btn"):
                        st.session_state.edited_content = edited_content
                        st.success("Content changes saved successfully!")

                # Content Expansion Feature
                with st.expander("✨ Expand Selected Content"):
                    selected_text = st.text_area(
                        "Paste the section you want to expand",
                        height=200,
                        key="expansion_input"
                    )
                    if st.button("Expand Content", key="expand_content_btn") and selected_text:
                        with st.spinner("Expanding content..."):
                            expanded_text = expand_content(
                                selected_text,
                                st.session_state.primary_keyword
                            )
                            st.text_area(
                                "Expanded Content (Copy and paste back to main editor)",
                                value=expanded_text,
                                height=300,
                                key="expanded_output"
                            )

                # SEO Metadata Editor
                with st.expander("📊 SEO Metadata Editor"):
                    edited_title = st.text_input(
                        "SEO Title",
                        value=st.session_state.edited_seo.get('title', ''),
                        max_chars=60,
                        help="Maximum 60 characters",
                        key="seo_title_input"
                    )
                    st.caption(f"Character count: {len(edited_title)}/60")

                    edited_meta = st.text_area(
                        "Meta Description",
                        value=st.session_state.edited_seo.get('meta_description', ''),
                        max_chars=160,
                        help="Maximum 160 characters",
                        key="seo_meta_input"
                    )
                    st.caption(f"Character count: {len(edited_meta)}/160")

                    edited_schema = st.text_area(
                        "Schema Markup",
                        value=json.dumps(st.session_state.edited_seo.get('schema', {}), indent=2),
                        height=300,
                        key="seo_schema_input"
                    )

                    if st.button("Save SEO Changes", key="save_seo_btn"):
                        try:
                            schema_dict = json.loads(edited_schema)
                            st.session_state.edited_seo = {
                                'title': edited_title,
                                'meta_description': edited_meta,
                                'schema': schema_dict
                            }
                            st.success("SEO changes saved successfully!")
                        except json.JSONDecodeError:
                            st.error("Invalid JSON in schema markup")

                # Save content changes
                if st.button("Save Content Changes"):
                    st.session_state.edited_content = edited_content
                    st.success("Content changes saved successfully!")

                # Template Analysis
                with st.expander("🔍 Template Analysis"):
                    st.write(st.session_state.generation_content['template_analysis'])

                # Download Section
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

                # Prepare files using edited content
                doc = create_word_doc(st.session_state.edited_content)
                word_buffer = BytesIO()
                doc.save(word_buffer)
                word_buffer.seek(0)

                csv_buffer = BytesIO()
                csv_buffer.write(save_to_csv({
                    'title': st.session_state.edited_seo['title'],
                    'meta_description': st.session_state.edited_seo['meta_description'],
                    'schema': json.dumps(st.session_state.edited_seo['schema'])
                }))
                csv_buffer.seek(0)

                # Create ZIP
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr(f'content_{timestamp}.docx', word_buffer.getvalue())
                    zf.writestr(f'content_{timestamp}.md',
                                st.session_state.edited_content.encode('utf-8'))
                    zf.writestr(f'metadata_{timestamp}.csv', csv_buffer.getvalue())
                zip_buffer.seek(0)

                # Download buttons
                col1, col2 = st.columns([2, 1])
                with col1:
                    st.download_button(
                        "📦 Download All Files (ZIP)",
                        zip_buffer.getvalue(),
                        f"generated_content_{timestamp}.zip",
                        "application/zip",
                        use_container_width=True
                    )

                with col2:
                    with st.expander("Download Individual Files"):
                        st.download_button("📥 Download Word",
                                           word_buffer.getvalue(),
                                           f"content_{timestamp}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        st.download_button("📥 Download Markdown",
                                           st.session_state.edited_content,
                                           f"content_{timestamp}.md",
                                           "text/markdown")
                        st.download_button("📥 Download Metadata",
                                           csv_buffer.getvalue(),
                                           f"metadata_{timestamp}.csv",
                                           "text/csv")
                
                with st.expander("🌐 Generate HTML"):
                    # Check for templates at the start
                    st.write("Checking for template files...")
                    available_templates = check_templates()
                    
                    if not available_templates:
                        st.error("No template files found! Please ensure template files (e.g., astoria-template.html) are in the correct directory.")
                        return
                        
                    try:
                        page_type = st.selectbox(
                            "Select Page Type",
                            options=['service page'],
                            key="page_type_select"
                        )
                        
                        # Show available templates
                        st.info(f"Available templates: {', '.join(available_templates)}")
                        
                        site_name = st.text_input(
                            "Enter Site Name",
                            key="site_name_input",
                            help=f"Available templates: {', '.join([t.replace('-template.html', '') for t in available_templates])}"
                        )
                        
                        if site_name:
                            template_file = f"{site_name.lower()}-template.html"
                            if not template_file in available_templates:
                                st.error(f"Template file '{template_file}' not found. Please choose from available templates.")
                                return
                        
                        if st.button("Generate HTML", key="generate_html_btn"):
                            if page_type and site_name:
                                try:
                                    debug_print("Creating CSV data...")
                                    csv_data = {
                                        'TITLE': [st.session_state.edited_seo.get('title', '')],
                                        'META_DESC': [st.session_state.edited_seo.get('meta_description', '')],
                                        'FAQ_SCHEMA': [json.dumps(st.session_state.edited_seo.get('schema', {}))],
                                        'CONTENT': [st.session_state.edited_content],
                                        'H2': [site_name]
                                    }
                                    
                                    # Create and save temporary CSV file
                                    df = pd.DataFrame(csv_data)
                                    temp_csv_path = os.path.join(os.getcwd(), 'temp.csv')
                                    df.to_csv(temp_csv_path, index=False)
                                    debug_print(f"Saved CSV to: {temp_csv_path}")
                                    debug_print(f"CSV exists: {os.path.exists(temp_csv_path)}")
                                    debug_print(f"CSV content preview: \n{df.head()}")
                                    
                                    debug_print("Calling generate_filled_html...")
                                    template_name = site_name.lower()
                                    generate_filled_html(temp_csv_path, template_name)
                                    
                                    output_filename = f"{template_name}.html"
                                    if os.path.exists(output_filename):
                                        with open(output_filename, 'r', encoding='utf-8') as f:
                                            html_content = f.read()
                                        
                                        st.download_button(
                                            "📥 Download Generated HTML",
                                            html_content,
                                            f"{template_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                                            "text/html"
                                        )
                                        
                                        # Clean up
                                        os.remove(output_filename)
                                        os.remove(temp_csv_path)
                                        st.success("HTML generated successfully!")
                                    else:
                                        st.error(f"Output file {output_filename} was not created")
                                        
                                except Exception as e:
                                    st.error(f"Error generating HTML: {str(e)}")
                                    st.write(traceback.format_exc())
                                    if os.path.exists(temp_csv_path):
                                        os.remove(temp_csv_path)
                            else:
                                st.warning("Please select a page type and enter a site name.")
                    except Exception as e:
                        st.error(f"Error in HTML generation section: {str(e)}")
                        st.write(traceback.format_exc())

if __name__ == "__main__":
    main()
